import argparse
import json
import re
from pathlib import Path

# Nếu bạn dùng .docx, cần cài python-docx:
try:
    from docx import Document  # type: ignore
except Exception:
    Document = None  # Cho phép chạy với .txt mà không cần cài


def read_txt_as_lines(txt_path: Path, encoding: str = "utf-8"):
    """Đọc file .txt -> list dòng (bỏ dòng rỗng đầu/cuối, giữ thứ tự)."""
    with txt_path.open("r", encoding=encoding, errors="replace") as f:
        # strip() từng dòng; bỏ dòng trống hoàn toàn
        lines = [line.strip() for line in f]
    # Không loại bỏ dòng trống giữa answer để tránh mất ngắt đoạn.
    # Thay vì loại bỏ tất cả, ta chỉ loại bỏ rỗng ở đầu file.
    while lines and not lines[0]:
        lines.pop(0)
    while lines and not lines[-1]:
        lines.pop()
    return lines


def read_docx_as_lines(docx_path: Path):
    """Đọc .docx -> list dòng. Gộp paragraphs và cả nội dung trong bảng."""
    if Document is None:
        raise RuntimeError(
            "Thiếu thư viện 'python-docx'. Hãy cài: pip install python-docx"
        )

    doc = Document(str(docx_path))
    lines = []

    # Paragraphs
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        # Giữ cả dòng trống (để bảo toàn ngắt đoạn trong answer),
        # nhưng bỏ bớt các đoạn rác hoàn toàn trống ở đầu/cuối sau này.
        lines.append(t)

    # Tables (nếu có)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = "\n".join(par.text for par in cell.paragraphs)
                for ln in cell_text.splitlines():
                    lines.append(ln.strip())

    # Cắt rỗng ở đầu/cuối
    while lines and not lines[0]:
        lines.pop(0)
    while lines and not lines[-1]:
        lines.pop()
    return lines


def parse_qa_from_lines(lines):
    """
    Parse định dạng:
    n. <câu hỏi ...?>   # có thể có nhiều dấu ? trong câu hỏi (ví dụ trong ngoặc đơn)
    <các dòng trả lời...> (có thể nhiều đoạn, hoặc tiếp tục trên cùng dòng)

    Logic cải thiện:
    - Nhận diện toàn bộ câu hỏi cho đến dấu ? cuối cùng
    - Bỏ qua các dấu ? ở giữa câu (như trong ví dụ: (ví dụ: ...? ...)
    - Phần sau dấu ? cuối cùng là câu trả lời

    Trả về list dict với các trường: "question", "answer" (không chứa "id").
    """
    # Ghép tất cả dòng thành một chuỗi lớn, giữ nguyên ngắt dòng
    full_text = "\n".join(lines)

    # Regex để tìm câu hỏi: số. văn bản kết thúc bằng ? hoặc :
    # Cải thiện để handle trường hợp có phần giải thích sau dấu ?
    # Ưu tiên match dấu ? ở cuối câu, bỏ qua dấu ? trong ngoặc đơn
    # Chỉ match ở đầu dòng
    question_pattern = r"^(\d+)\.\s+([^?]*?)\?+(?=\s|$|\n)"

    # Tìm tất cả vị trí câu hỏi
    matches = list(re.finditer(question_pattern, full_text, re.MULTILINE))

    items = []
    for i, match in enumerate(matches):
        cur_id = int(match.group(1))
        q_text = match.group(2).strip() + "?"  # thêm dấu ? vào cuối câu hỏi

        # Xác định vị trí bắt đầu answer: ngay sau match
        start_answer = match.end()

        # Vị trí kết thúc answer: trước câu hỏi tiếp theo hoặc cuối text
        if i + 1 < len(matches):
            end_answer = matches[i + 1].start()
        else:
            end_answer = len(full_text)

        # Lấy answer, loại bỏ khoảng trắng thừa
        answer_text = full_text[start_answer:end_answer].strip()

        items.append((cur_id, {"question": q_text, "answer": answer_text}))

    return [obj for _, obj in items]


def save_to_jsonl(qa_items, out_path: Path):
    """Lưu list các dict thành JSONL (mỗi dòng một object)."""
    # Ghi từng item trên một dòng (UTF-8, không escape Unicode)
    with out_path.open("w", encoding="utf-8") as f:
        for item in qa_items:
            f.write(json.dumps(item, ensure_ascii=False) + "\n")
    return out_path


def convert(
    input_path: str, output_path: str | None = None, txt_encoding: str = "utf-8"
):
    in_path = Path(input_path)
    if not in_path.exists():
        raise FileNotFoundError(f"Không tìm thấy file: {in_path}")
    if output_path is None:
        out_path = in_path.with_suffix(".jsonl")
    else:
        out_path = Path(output_path)

    ext = in_path.suffix.lower()
    if ext == ".txt":
        lines = read_txt_as_lines(in_path, encoding=txt_encoding)
    elif ext == ".docx":
        lines = read_docx_as_lines(in_path)
    else:
        raise ValueError("Định dạng không hỗ trợ. Vui lòng dùng .txt hoặc .docx")

    qa_items = parse_qa_from_lines(lines)
    save_to_jsonl(qa_items, out_path)
    return qa_items, out_path


def main(argv=None):
    parser = argparse.ArgumentParser(
        description="Trích xuất Q&A (n. ...? hoặc :) từ .txt/.docx và lưu JSONL (mỗi dòng một object)."
    )
    parser.add_argument(
        "--in", dest="inp", required=True, help="Đường dẫn file .txt hoặc .docx"
    )
    parser.add_argument(
        "--out",
        dest="outp",
        default=None,
        help="Đường dẫn file JSONL đầu ra (mặc định cùng tên .jsonl)",
    )
    parser.add_argument(
        "--encoding",
        dest="enc",
        default="utf-8",
        help="Mã hóa khi đọc .txt (mặc định utf-8)",
    )
    args = parser.parse_args(argv)

    qa, outp = convert(args.inp, args.outp, args.enc)
    print(f"Đã trích xuất {len(qa)} mục Q&A -> {outp}")


if __name__ == "__main__":
    main()
